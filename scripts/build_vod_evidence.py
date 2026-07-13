from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "开发者声音VOD建议材料.docx"

BLUE = RGBColor(46, 116, 181)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(91, 110, 133)

def font(run, size=11, color=None, bold=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if bold is not None: run.bold = bold

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def set_width(cell, width):
    tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.find(qn("w:tcW"))
    if tcW is None: tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width)); tcW.set(qn("w:type"), "dxa")

def cell_text(cell, text, bold=False, color=None):
    cell.text = ""; p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text); font(r, 10.5, color, bold); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc = Document(); section = doc.sections[0]
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(.492)
styles = doc.styles
styles["Normal"].font.name = "Calibri"; styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(11); styles["Normal"].paragraph_format.space_after = Pt(6); styles["Normal"].paragraph_format.line_spacing = 1.1

header = section.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = header.add_run("Developer Retention Radar | Competition Evidence"); font(r, 9, MUTED)
footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("开发者留存雷达 · 华为云赛事交付佐证材料"); font(r, 9, MUTED)

p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
r = p.add_run("开发者声音（VOD）建议材料"); font(r, 24, NAVY, True)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(16)
r = p.add_run("聚焦 Flexus 应用服务器 L 实例容器化部署体验优化"); font(r, 12.5, MUTED)

table = doc.add_table(rows=4, cols=2); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = "Table Grid"
for row, (label, value) in zip(table.rows, [
    ("关联产品", "华为云 Flexus 云服务 / Flexus 应用服务器 L 实例"),
    ("建议类型", "用户体验与功能优化建议"),
    ("项目名称", "开发者留存雷达（Developer Retention Radar）"),
    ("材料状态", "建议稿，随竞赛作品压缩包提交"),
]):
    set_width(row.cells[0], 2700); set_width(row.cells[1], 6660); shade(row.cells[0], "E8EEF5")
    cell_text(row.cells[0], label, True, NAVY); cell_text(row.cells[1], value)

def heading(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); font(r, 14, BLUE, True)

def para(text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(7); p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text); font(r, 11)

heading("一、使用场景与问题记录")
para("在华为云 Flexus 应用服务器 L 实例上使用 Docker 可视化 Portainer 镜像部署“开发者留存雷达”时，需要通过 Docker Compose 将 Web 服务映射到 80 端口。实际部署发现宿主机 Nginx 默认占用了 80 端口，Compose 启动时返回 address already in use，需要自行排查并停止宿主机服务后才能完成上线。")

heading("二、优化建议")
para("建议在 Flexus 镜像购买页、实例首次登录指引和 Portainer 控制台中明确展示宿主机默认端口占用情况，并提供端口清单、端口冲突检测和“一键释放 80 端口”入口。对于常见的 Docker Compose Web 应用，可增加“标准 Web 部署”预检提示，提前告知端口映射冲突与推荐处理方式。")

heading("三、预期价值")
for item in [
    "减少初次使用 Flexus 容器镜像的排障时间，降低新手部署门槛。",
    "避免用户因端口冲突误判 Docker、Compose 或应用配置异常。",
    "提升从购买实例到公网访问 Web 应用的端到端体验。",
]:
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(item); font(r, 11)

heading("四、项目验证信息")
para("本建议来自真实竞赛项目部署过程。项目采用 React、Node.js、PostgreSQL 与 Docker Compose，已部署在华为云 Flexus 应用服务器 L 实例，并完成公网访问、签到、漏斗统计、作品提交与导出等流程验证。")
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(5)
r = p.add_run("公网访问："); font(r, 11, NAVY, True); r = p.add_run("http://123.60.220.148"); font(r, 11, BLUE)

OUT.parent.mkdir(exist_ok=True)
doc.core_properties.title = "开发者声音VOD建议材料"
doc.core_properties.author = "Developer Retention Radar"
doc.save(OUT)
print(OUT)
